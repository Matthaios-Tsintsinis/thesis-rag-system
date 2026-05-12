"""Document parsing.

PDF parser: **Docling** (single source of truth for PDF across all
systems). The pipeline design uses Docling because M7's structural
axis depends on its section metadata; running a second PDF parser
elsewhere would produce different text → different chunks → different
embeddings, confounding the M7-vs-baseline comparison.

Non-PDF parsers (docx/html/csv/json/xlsx/txt/md) ported from the
existing notebook unchanged. They synthesise a minimal "sections"
payload (single section, depth 0, title = filename) so downstream code
can treat every doc uniformly.

Cache invalidation: bump `PARSING_VERSION` whenever the parser identity
or per-parser output format changes. `parsing_identity()` is folded
into every retriever cache key via cache.compute_cache_key, so changes
here invalidate caches automatically.
"""

from __future__ import annotations

import json as _json
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable


# --- Identity for cache invalidation --------------------------------------

# Bump this whenever the parser identity or output schema changes so
# cached embeddings/indexes from older runs are invalidated cleanly.
PARSING_VERSION = "docling-v1"


def parsing_identity() -> dict:
    return {"pdf_parser": "docling", "parsing_version": PARSING_VERSION}


# --- Supported formats ----------------------------------------------------

SUPPORTED_EXTENSIONS = {
    ".pdf", ".txt", ".md",
    ".html", ".htm",
    ".docx", ".csv", ".json", ".xlsx",
}


@dataclass
class ParsedDocument:
    doc_id: str
    path: Path
    text: str
    metadata: dict = field(default_factory=dict)


# --- Helpers --------------------------------------------------------------


def safe_read_text(path: Path) -> str:
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def _fallback_sections(filename: str, text: str) -> list[dict]:
    """Minimal one-section payload for non-PDF docs.

    Per PIPELINE_DESIGN §3.3: non-PDF inputs get section_title =
    file_name, section_depth = 0, neighbours by chunk order. The
    neighbour links are assigned later when chunks are produced; here
    we just stub a single section spanning the whole document.
    """
    return [{
        "section_title": filename,
        "section_depth": 0,
        "section_path": [filename],
        "page_start": None,
        "page_end": None,
        "order_in_document": 0,
        "text": text,
    }]


# --- Per-format parsers ---------------------------------------------------


def parse_txt(path: Path) -> tuple[str, dict]:
    text = safe_read_text(path)
    return text, {"sections": _fallback_sections(Path(path).name, text)}


def parse_pdf(path: Path) -> tuple[str, dict]:
    """Parse a PDF with Docling.

    Returns:
        (text_with_page_markers, metadata)

    Metadata schema:
        {
          "sections": [
            {
              "section_title": str,
              "section_depth": int,         # heading level (1=H1, 2=H2, ...)
              "section_path": [str, ...],   # ancestor chain of section titles
              "page_start": int | None,
              "page_end":   int | None,
              "order_in_document": int,     # 0-indexed
              "text": str,                  # joined body text of section
            },
            ...
          ],
          "n_pages": int,
          "n_sections": int,
        }

    Text body interleaves `[PAGE n]` markers (one per new page) so
    `detect_page_refs` keeps working for citations. Section metadata is
    a separate payload — never embedded into chunk text per
    PIPELINE_DESIGN §3.3 ("Metadata is a control signal, not a search
    signal").
    """
    from docling.document_converter import DocumentConverter

    converter = DocumentConverter()
    result = converter.convert(str(path))
    doc = result.document

    text_parts: list[str] = []
    sections: list[dict] = []
    section_stack: list[tuple[int, str]] = []   # (depth, title)
    current_page: int | None = None
    order_idx = 0
    current_section: dict | None = None
    max_page = 0

    def _emit_marker(page_no: int | None) -> None:
        nonlocal current_page, max_page
        if page_no and page_no != current_page:
            text_parts.append(f"\n\n[PAGE {page_no}]\n")
            current_page = page_no
            if page_no > max_page:
                max_page = page_no

    def _close_section() -> None:
        nonlocal current_section
        if current_section is not None:
            sections.append(current_section)
            current_section = None

    # Docling 2.x API: doc.iterate_items() yields (item, level) tuples.
    # Older versions expose .body.items / .texts; we try the modern API
    # first and fall back to a text-only path if it isn't available.
    try:
        items_iter = doc.iterate_items()
    except AttributeError:
        items_iter = None

    if items_iter is not None:
        for item, level in items_iter:
            page_no: int | None = None
            prov = getattr(item, "prov", None)
            if prov:
                try:
                    page_no = int(prov[0].page_no)
                except (AttributeError, IndexError, TypeError):
                    page_no = None
            _emit_marker(page_no)

            label = str(getattr(item, "label", "") or "").lower()
            text = (getattr(item, "text", None) or "").strip()
            if not text:
                continue

            if label in {"section_header", "title", "heading"}:
                depth = max(1, int(level) or 1)
                while section_stack and section_stack[-1][0] >= depth:
                    section_stack.pop()
                section_stack.append((depth, text))

                _close_section()
                current_section = {
                    "section_title": text,
                    "section_depth": depth,
                    "section_path": [t for _, t in section_stack],
                    "page_start": page_no,
                    "page_end": page_no,
                    "order_in_document": order_idx,
                    "text": "",
                }
                order_idx += 1
                text_parts.append(text)
            else:
                text_parts.append(text)
                if current_section is None:
                    # Body text before any heading — make a synthetic intro section
                    current_section = {
                        "section_title": Path(path).name,
                        "section_depth": 0,
                        "section_path": [Path(path).name],
                        "page_start": page_no,
                        "page_end": page_no,
                        "order_in_document": order_idx,
                        "text": "",
                    }
                    order_idx += 1
                current_section["text"] = (
                    current_section["text"] + " " + text if current_section["text"]
                    else text
                )
                if page_no:
                    current_section["page_end"] = page_no
        _close_section()
    else:
        # Fallback: text only, no structure. Mark with synthetic single section.
        body = doc.export_to_text() if hasattr(doc, "export_to_text") else ""
        text_parts.append(body)
        sections = _fallback_sections(Path(path).name, body)

    if not sections:
        sections = _fallback_sections(Path(path).name, "\n".join(text_parts))

    return "\n".join(text_parts).strip(), {
        "sections": sections,
        "n_pages": max_page,
        "n_sections": len(sections),
    }


def parse_docx(path: Path) -> tuple[str, dict]:
    from docx import Document as DocxDocument

    d = DocxDocument(str(path))
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    text = "\n".join(paras)
    return text, {
        "sections": _fallback_sections(Path(path).name, text),
        "n_paragraphs": len(paras),
    }


def parse_html(path: Path) -> tuple[str, dict]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(safe_read_text(path), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.extract()
    text = soup.get_text("\n")
    return text, {"sections": _fallback_sections(Path(path).name, text)}


def parse_csv(path: Path) -> tuple[str, dict]:
    import pandas as pd

    df = pd.read_csv(path)
    sample = df.head(200).astype(str)
    txt = "\n".join(" | ".join(row) for row in sample.values.tolist())
    return txt, {
        "sections": _fallback_sections(Path(path).name, txt),
        "n_rows": int(df.shape[0]),
        "n_cols": int(df.shape[1]),
        "columns": list(map(str, df.columns)),
    }


def parse_json_file(path: Path) -> tuple[str, dict]:
    obj = _json.loads(Path(path).read_text(encoding="utf-8"))
    pretty = _json.dumps(obj, ensure_ascii=False, indent=2)
    return pretty, {
        "sections": _fallback_sections(Path(path).name, pretty),
        "root_type": type(obj).__name__,
    }


def parse_xlsx(path: Path) -> tuple[str, dict]:
    import pandas as pd

    xls = pd.ExcelFile(path)
    parts: list[str] = []
    sheet_shapes: dict[str, list[int]] = {}
    for sheet in xls.sheet_names[:10]:
        df = pd.read_excel(path, sheet_name=sheet)
        sheet_shapes[sheet] = [int(df.shape[0]), int(df.shape[1])]
        sample = df.head(100).astype(str)
        parts.append(f"\n[SHEET: {sheet}]\n")
        parts.append("\n".join(" | ".join(row) for row in sample.values.tolist()))
    text = "\n".join(parts)
    return text, {
        "sections": _fallback_sections(Path(path).name, text),
        "sheet_shapes": sheet_shapes,
        "sheet_names": xls.sheet_names,
    }


# --- Cleaning + dispatch --------------------------------------------------


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" ", " ", text)
    return text.strip()


def extract_text(path: Path) -> tuple[str, dict]:
    """Dispatch by extension. Returns (cleaned_text, metadata).

    Metadata always contains a `sections` list (per
    PIPELINE_DESIGN §3.3); for non-PDF inputs the list has one entry.
    """
    ext = Path(path).suffix.lower()
    metadata: dict = {"path": str(path), "ext": ext}
    if ext == ".pdf":
        text, extra = parse_pdf(path)
    elif ext == ".docx":
        text, extra = parse_docx(path)
    elif ext in {".html", ".htm"}:
        text, extra = parse_html(path)
    elif ext == ".csv":
        text, extra = parse_csv(path)
    elif ext == ".json":
        text, extra = parse_json_file(path)
    elif ext == ".xlsx":
        text, extra = parse_xlsx(path)
    elif ext in {".txt", ".md"}:
        text, extra = parse_txt(path)
    else:
        raise ValueError(f"Unsupported extension {ext!r} for {path}")
    metadata.update(extra)
    return clean_text(text), metadata


def parse_file(path: Path) -> str:
    """Back-compat: cleaned text only. Use extract_text for metadata."""
    text, _ = extract_text(path)
    return text


def detect_page_refs(chunk_text: str) -> list[int]:
    return [int(m) for m in re.findall(r"\[PAGE (\d+)\]", chunk_text)]


def list_files_recursive(root: Path) -> list[Path]:
    files: list[Path] = []
    for path, _, fnames in os.walk(str(root)):
        for f in fnames:
            full = Path(path) / f
            if full.suffix.lower() in SUPPORTED_EXTENSIONS:
                files.append(full)
    return sorted(files)


def walk_corpus(folder: Path, min_chars: int = 0) -> Iterable[ParsedDocument]:
    """Yield ParsedDocument for every supported file under `folder`.

    Docs whose cleaned text is below `min_chars` are skipped.
    """
    folder = Path(folder)
    for path in list_files_recursive(folder):
        try:
            text, meta = extract_text(path)
        except Exception as e:
            print(f"[parsing] skip {path}: {type(e).__name__}: {e}")
            continue
        if len(text) < min_chars:
            continue
        doc_id = str(path.relative_to(folder)).replace("\\", "/")
        yield ParsedDocument(doc_id=doc_id, path=path, text=text, metadata=meta)
